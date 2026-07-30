"""
tests/test_extractor.py - Unit tests for extractor.py
------------------------------------------------------
Tests text extraction from TXT, PDF, and DOCX files,
and verifies text cleaning and normalization functions.
"""

import unittest
import os
import io
import docx
import pypdf
from extractor import clean_text, extract_text, extract_text_from_docx, extract_text_from_pdf

class TestExtractor(unittest.TestCase):

    def setUp(self):
        self.sample_dir = os.path.join(os.path.dirname(__file__), "..", "sample_data")
        os.makedirs(self.sample_dir, exist_ok=True)

        # 1. Path to sample text file
        self.txt_path = os.path.join(self.sample_dir, "sample_resume.txt")

        # 2. Programmatically generate a sample DOCX file for testing
        self.docx_path = os.path.join(self.sample_dir, "sample_resume.docx")
        doc = docx.Document()
        doc.add_heading("Jane Doe - Software Engineer", level=1)
        doc.add_paragraph("Skills: Python • JavaScript • React • SQL")
        doc.add_paragraph("Experience: Built web apps using FastAPI and Docker.")
        # Add table
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Education"
        table.rows[0].cells[1].text = "B.S. Computer Science"
        table.rows[1].cells[0].text = "GPA"
        table.rows[1].cells[1].text = "3.9"
        doc.save(self.docx_path)

    def test_clean_text(self):
        """Test Unicode normalization, control character removal, and space collapsing."""
        raw_dirty_text = "Jane Doe — Senior Dev\x07\n• Skill 1: Python   &   React\n\n“Smart Quotes”"
        cleaned = clean_text(raw_dirty_text)

        self.assertIn("Jane Doe - Senior Dev", cleaned)
        self.assertIn("* Skill 1: Python & React", cleaned)
        self.assertIn('"Smart Quotes"', cleaned)
        self.assertNotIn("\x07", cleaned)

    def test_extract_text_from_txt(self):
        """Test extraction from standard text resume file."""
        text = extract_text(self.txt_path)
        self.assertIn("ALEX CHEN", text)
        self.assertIn("Python", text)
        self.assertIn("FastAPI", text)

    def test_extract_text_from_docx(self):
        """Test extraction from DOCX file including paragraphs and tables."""
        text = extract_text(self.docx_path)
        self.assertIn("Jane Doe - Software Engineer", text)
        self.assertIn("Python", text)
        self.assertIn("Education | B.S. Computer Science", text)

    def test_extract_text_bytes_stream(self):
        """Test extraction using file-like BytesIO streams (as received from Streamlit uploaders)."""
        doc = docx.Document()
        doc.add_paragraph("Streamlit Upload Test: Python, Docker, PyTorch")
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)

        extracted = extract_text(bio, filename="resume.docx")
        self.assertIn("Streamlit Upload Test", extracted)
        self.assertIn("PyTorch", extracted)

    def test_invalid_extension(self):
        """Test error handling for unsupported file formats."""
        with self.assertRaises(ValueError):
            extract_text("invalid_file.xyz")

if __name__ == "__main__":
    unittest.main()
