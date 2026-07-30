"""
extractor.py - Resume Text Extraction Engine
---------------------------------------------
This module handles extracting text from PDF and DOCX resume files,
and performs text cleaning and normalization to prepare documents for downstream NLP task.
"""

import io
import re
import os
import pypdf
import docx

def clean_text(text: str) -> str:
    """
    Cleans and normalizes extracted text.

    - Normalizes Unicode punctuation (curly quotes, dashes, bullet points)
    - Removes non-printable ASCII control characters
    - Normalizes excessive whitespace and trailing spaces per line
    - Preserves sentence structure and paragraph breaks

    Args:
        text (str): Raw text extracted from document.

    Returns:
        str: Cleaned and normalized text string.
    """
    if not text:
        return ""

    # Replace common Unicode symbols and punctuation
    replacements = {
        '“': '"', '”': '"', "‘": "'", "’": "'",
        '–': '-', '—': '-', '…': '...',
        '•': '*', '▪': '*', '▸': '*', '►': '*', '●': '*', '◦': '*'
    }
    for orig, repl in replacements.items():
        text = text.replace(orig, repl)

    # Remove non-printable control characters (except standard newlines and tabs)
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', ' ', text)

    # Normalize multiple horizontal spaces to single space
    text = re.sub(r'[ \t]+', ' ', text)

    # Split lines, strip leading/trailing whitespace per line, ignore empty lines
    lines = [line.strip() for line in text.splitlines()]
    
    # Rejoin with clean newlines, removing redundant consecutive blank lines
    cleaned_lines = []
    prev_empty = False
    for line in lines:
        if line:
            cleaned_lines.append(line)
            prev_empty = False
        elif not prev_empty:
            cleaned_lines.append("")
            prev_empty = True

    return "\n".join(cleaned_lines).strip()


def extract_text_from_pdf(file_source) -> str:
    """
    Extracts raw text from a PDF file path or file-like byte stream.

    Args:
        file_source (str | BytesIO | UploadedFile): Path to PDF or binary stream.

    Returns:
        str: Extracted raw text.
    """
    raw_text = ""
    try:
        if isinstance(file_source, (str, os.PathLike)):
            reader = pypdf.PdfReader(file_source)
        else:
            # Handle Streamlit UploadedFile or BytesIO
            file_bytes = io.BytesIO(file_source.read()) if hasattr(file_source, 'read') else file_source
            reader = pypdf.PdfReader(file_bytes)

        for page_idx, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                raw_text += page_text + "\n"
    except Exception as e:
        raise ValueError(f"Failed to parse PDF document: {str(e)}")

    return raw_text


def extract_text_from_docx(file_source) -> str:
    """
    Extracts raw text from a DOCX file path or file-like byte stream.
    Includes text from paragraphs and table cells.

    Args:
        file_source (str | BytesIO | UploadedFile): Path to DOCX or binary stream.

    Returns:
        str: Extracted raw text.
    """
    raw_text = []
    try:
        if isinstance(file_source, (str, os.PathLike)):
            doc = docx.Document(file_source)
        else:
            # Handle Streamlit UploadedFile or BytesIO
            file_bytes = io.BytesIO(file_source.read()) if hasattr(file_source, 'read') else file_source
            doc = docx.Document(file_bytes)

        # Extract text from standard paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                raw_text.append(paragraph.text)

        # Extract text from tables (common in resume templates)
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    raw_text.append(" | ".join(row_data))

    except Exception as e:
        raise ValueError(f"Failed to parse DOCX document: {str(e)}")

    return "\n".join(raw_text)


def extract_text(file_source, filename: str = None) -> str:
    """
    Main dispatcher function to extract and clean text from PDF or DOCX documents.

    Args:
        file_source (str | BytesIO | UploadedFile): File path or uploaded file object.
        filename (str, optional): Name of the file to determine extension if file_source is a stream.

    Returns:
        str: Cleaned and normalized text.
    """
    # Determine filename from string path if not explicitly provided
    if filename is None and isinstance(file_source, str):
        filename = file_source
    elif filename is None and hasattr(file_source, 'name'):
        filename = file_source.name

    if not filename:
        raise ValueError("Filename or file path must be provided to determine document format.")

    ext = os.path.splitext(filename)[1].lower()

    if ext == '.pdf':
        raw_text = extract_text_from_pdf(file_source)
    elif ext in ['.docx', '.doc']:
        raw_text = extract_text_from_docx(file_source)
    elif ext == '.txt':
        if isinstance(file_source, str):
            with open(file_source, 'r', encoding='utf-8', errors='ignore') as f:
                raw_text = f.read()
        else:
            file_bytes = file_source.read()
            raw_text = file_bytes.decode('utf-8', errors='ignore') if isinstance(file_bytes, bytes) else str(file_bytes)
    else:
        raise ValueError(f"Unsupported file format '{ext}'. Only PDF, DOCX, and TXT are supported.")

    return clean_text(raw_text)
